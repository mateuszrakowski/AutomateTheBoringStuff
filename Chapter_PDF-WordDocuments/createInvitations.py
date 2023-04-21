"""
Creates party invitations in .docx format based on guests list provided in txt file.

Usage:
python createInvitations.py <txt file>

Arguments:
txt file (str): Path/name of txt file with list of guests.

Returns:
Document in docx format with invitations.
"""

import argparse
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Pt


def checkExtension(filename):
    """
    Check if provided file is txt format.

    Returns:
    str: The filename entered by user.
    """
    if not filename.endswith(".txt"):
        raise argparse.ArgumentTypeError("File must have .txt extension.")
    return filename


def checkFile():
    """
    Prompts the user to enter a text filename or path.

    Returns:
    str: The filename entered by the user.
    """
    parser = argparse.ArgumentParser(description="Create invitations in docx document.")
    parser.add_argument(
        "filename",
        help="Name or path of text file with guests names.",
        type=checkExtension,
    )
    args = parser.parse_args()
    return args.filename


def getGuests():
    """
    Iterate through every line in provided filename and store it in list.

    Returns:
    list: A list of guests.
    """
    txtFile = checkFile()
    p = Path(txtFile)

    listOfGuests = []

    try:
        with open(p.absolute(), "r") as f:
            for guest in f:
                listOfGuests.append(guest.strip())
    except:
        print("Couldn't open text file.")

    return listOfGuests


def createInvitation(guests):
    """
    Creates invitations based on guests list in one docx document.

    Returns:
    Docx document with invitations.
    """
    document = Document()
    style = document.styles["Normal"]
    font = style.font
    font.size = Pt(20)

    for guest in guests:
        paragraph = document.add_paragraph()
        paragraph.style = document.styles["Normal"]
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = paragraph.add_run("It would be a pleasure to have the company of")
        run.italic = True
        run.add_break()

        run = paragraph.add_run(guest)
        run.bold = True
        run.add_break()

        run = paragraph.add_run("at 1101 Memory Lane on the Evening of")
        run.italic = True
        run.add_break()

        run = paragraph.add_run("April 1st")
        run.add_break()

        run = paragraph.add_run("at 7 o'clock")
        run.italic = True
        run.add_break(WD_BREAK.PAGE)

    document.save("Invitations.docx")
    print("Done!")


if __name__ == "__main__":
    createInvitation(getGuests())
